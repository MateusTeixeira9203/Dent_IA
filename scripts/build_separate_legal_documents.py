from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

import build_legal_dossier as base


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "juridico" / "documentos-separados"
VERSION = "Versao 0.1 | 3 de setembro de 2026"


MAIN_DOCUMENTS = [
    ("01-termos-de-uso-e-contrato-saas.docx", "Termos de Uso e Contrato de Licenca SaaS", base.build_terms),
    ("02-aviso-de-privacidade-dos-usuarios.docx", "Aviso de Privacidade de Usuarios da Plataforma", base.build_user_privacy),
    ("03-acordo-de-tratamento-de-dados-dpa.docx", "Acordo de Tratamento de Dados Pessoais (DPA)", base.build_dpa),
    ("04-aviso-de-privacidade-do-paciente.docx", "Aviso de Privacidade do Paciente", base.build_patient_privacy),
    ("05-contrato-de-prestacao-de-servicos-odontologicos.docx", "Contrato de Prestacao de Servicos Odontologicos", base.build_service_contract),
    ("06-plano-de-tratamento-e-aceite-de-orcamento.docx", "Plano de Tratamento e Aceite de Orcamento", base.build_budget),
    ("07-tcle-nucleo-geral.docx", "TCLE - Nucleo Geral para Tratamento Odontologico", base.build_tcle),
    ("09-termo-de-recusa-informada.docx", "Termo de Recusa Informada, Interrupcao ou Descontinuidade", base.build_refusal),
    ("10-registro-de-conclusao-intercorrencias-e-orientacoes.docx", "Registro de Conclusao, Intercorrencias e Orientacoes", base.build_completion),
    ("11-autorizacao-audio-voz-imagem-transcricao.docx", "Autorizacao Especifica para Audio, Voz, Imagem e Transcricao", base.build_recording),
    ("12-politica-interna-geracao-assinatura.docx", "Politica Interna de Geracao e Assinatura", base.build_operational_rules),
]


MODULES = [
    ("A", "Cirurgia oral e extracoes", "08a-modulo-risco-cirurgia-oral-e-extracoes.docx"),
    ("B", "Implante, enxerto osseo e levantamento de seio", "08b-modulo-risco-implante-enxerto-e-seio.docx"),
    ("C", "Tratamento endodontico e retratamento", "08c-modulo-risco-endodontia-e-retratamento.docx"),
    ("D", "Cirurgia e terapia periodontal", "08d-modulo-risco-periodontia.docx"),
    ("E", "Ortodontia e alinhadores", "08e-modulo-risco-ortodontia-e-alinhadores.docx"),
    ("F", "Protese, facetas, coroas e procedimentos esteticos", "08f-modulo-risco-protese-e-estetica.docx"),
    ("G", "Anestesia local; sedacao exige termo proprio", "08g-modulo-risco-anestesia-local.docx"),
]


def clear_first_page_break(doc: Document) -> None:
    if doc.paragraphs:
        doc.paragraphs[0].paragraph_format.page_break_before = False


def set_metadata(doc: Document, title: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = "Minuta juridica para revisao profissional"
    doc.core_properties.author = "Odonto.IA"
    doc.core_properties.keywords = "odontologia, contratos, TCLE, LGPD, SaaS, prontuario"


def finish_document(doc: Document, title: str, path: Path) -> None:
    clear_first_page_break(doc)
    base.setup_header_footer(doc)
    set_metadata(doc, title)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_main_documents() -> list[Path]:
    outputs: list[Path] = []
    for filename, title, builder in MAIN_DOCUMENTS:
        doc = Document()
        base.setup_styles(doc)
        builder(doc)
        path = OUTPUT_DIR / filename
        finish_document(doc, title, path)
        outputs.append(path)
    return outputs


def module_source_ranges() -> dict[str, list]:
    source = Document()
    base.setup_styles(source)
    base.build_risk_modules(source)
    paragraphs = list(source.paragraphs)
    starts: dict[str, int] = {}
    for index, paragraph in enumerate(paragraphs):
        for letter, title, _ in MODULES:
            if paragraph.text == f"Modulo {letter} - {title}":
                starts[letter] = index

    if len(starts) != len(MODULES):
        missing = sorted({letter for letter, _, _ in MODULES} - set(starts))
        raise RuntimeError(f"Modulos nao encontrados no dossie: {missing}")

    ranges: dict[str, list] = {}
    ordered = [(letter, starts[letter]) for letter, _, _ in MODULES]
    for position, (letter, start) in enumerate(ordered):
        end = ordered[position + 1][1] if position + 1 < len(ordered) else len(paragraphs)
        ranges[letter] = paragraphs[start:end]
    return ranges


def add_module_header(doc: Document, letter: str, title: str) -> None:
    base.add_para(doc, f"MINUTA 8{letter}", bold=True, color=base.GOLD, size=10, after=4, keep=True)
    base.add_para(doc, f"MODULO DE RISCO - {title.upper()}", bold=True, color=base.INK, size=20, after=4, keep=True)
    base.add_para(doc, "Anexo individualizavel ao TCLE - Nucleo Geral para Tratamento Odontologico", italic=True, color=base.MUTED, size=10.5, after=12)
    base.add_callout(
        doc,
        "STATUS",
        "Rascunho para revisao por advogado brasileiro e validacao clinica por especialista. Nao publicar nem coletar assinatura antes do preenchimento dos campos e das aprovacoes necessarias.",
        fill=base.LIGHT_GOLD,
        color=base.RED,
    )
    base.add_callout(
        doc,
        "INTEGRACAO OBRIGATORIA",
        "Este modulo nao substitui o TCLE geral, a conversa clinica nem o registro no prontuario. Deve ser anexado ao TCLE, ajustado ao procedimento e ao paciente, e assinado no mesmo conjunto documental. Itens irrelevantes devem ser removidos e riscos materiais ausentes devem ser acrescentados.",
        fill=base.LIGHT_BLUE,
        color=base.INK,
    )
    base.add_para(doc, "Clinica: [[NOME/CNPJ/EPAO]] | Paciente: [[NOME/CPF/NASCIMENTO]]")
    base.add_para(doc, "Cirurgiao-dentista: [[NOME/CRO-UF]] | Procedimento/dente/regiao/lateralidade: [[DESCRICAO INEQUIVOCA]]")


def build_module_documents() -> list[Path]:
    source_ranges = module_source_ranges()
    outputs: list[Path] = []
    for letter, title, filename in MODULES:
        doc = Document()
        base.setup_styles(doc)
        add_module_header(doc, letter, title)
        for paragraph in source_ranges[letter]:
            doc._body._element.insert(-1, deepcopy(paragraph._p))
        base.add_callout(
            doc,
            "DECLARACAO PARA VALIDACAO",
            "O paciente recebeu explicacao individualizada dos riscos selecionados, de sua relevancia para o caso, das alternativas, dos cuidados e dos sinais de alarme. Este anexo integra o TCLE geral e nao representa renuncia a direitos.",
            fill=base.LIGHT_BLUE,
        )
        base.add_signature_block(
            doc,
            [
                "Paciente/Responsavel: [[NOME/CPF/QUALIDADE]]",
                "Cirurgiao-dentista: [[NOME/CRO-UF]]",
            ],
        )
        path = OUTPUT_DIR / filename
        finish_document(doc, f"Modulo de risco - {title}", path)
        outputs.append(path)
    return outputs


def build_guide(document_paths: list[Path]) -> Path:
    doc = Document()
    base.setup_styles(doc)
    base.add_para(doc, "ODONTO.IA", bold=True, color=base.GOLD, size=10, after=18)
    base.add_para(doc, "Guia de envio ao advogado", bold=True, color=base.INK, size=26, after=4)
    base.add_para(doc, "Pacote de minutas juridicas e documentos clinicos separados", color=base.DARK_BLUE, size=14, after=12)
    base.add_para(doc, VERSION, bold=True, color=base.MUTED, size=10.5, after=18)
    base.add_callout(
        doc,
        "NAO PUBLICAR",
        "As minutas contem campos [[PREENCHER]] e decisoes em aberto. O pacote serve para revisao juridica, validacao clinica e conferencia tecnica; nao esta pronto para coleta de aceite ou assinatura.",
        fill=base.LIGHT_GOLD,
        color=base.RED,
    )
    base.add_heading(doc, "1. Como o pacote esta organizado", 1)
    base.add_para(doc, "Os documentos foram separados para que cada relacao juridica e cada momento do atendimento tenham finalidade, signatarios e prova proprios. A separacao nao elimina as vinculacoes indicadas em cada minuta.")
    base.add_list_item(doc, "Plataforma: Termos de Uso, Aviso de Privacidade de Usuarios e DPA.")
    base.add_list_item(doc, "Clinica e paciente: Aviso de Privacidade do Paciente, contrato odontologico, plano/orcamento, TCLE, recusa e registros posteriores.")
    base.add_list_item(doc, "Procedimentos: cada modulo de risco e um anexo ao TCLE geral e exige individualizacao e validacao da especialidade.")
    base.add_list_item(doc, "Operacao interna: politica de geracao, revisao, assinatura, versionamento e conservacao da prova.")

    base.add_heading(doc, "2. Arquivos entregues", 1)
    rows: list[list[str]] = []
    for path in document_paths:
        category = "Clinica/paciente"
        if path.name.startswith(("01", "02", "03")):
            category = "Plataforma"
        if path.name.startswith("08"):
            category = "Anexo de risco"
        if path.name.startswith("12"):
            category = "Politica interna"
        rows.append([path.name.split("-")[0].upper(), path.stem.split("-", 1)[1].replace("-", " ").title(), category])
    base.add_table(doc, ["Codigo", "Documento", "Camada"], rows, [1200, 5360, 2800])

    base.add_heading(doc, "3. Pontos prioritarios para o advogado", 1)
    priorities = [
        "Qualificacao da empresa, da clinica, do contratante, do representante e dos profissionais; poderes para contratar e assinar.",
        "Incidencia do CDC, limites de responsabilidade, cobranca, renovacao, cancelamento, exportacao, continuidade e foro.",
        "Papeis de controlador e operador, bases legais por finalidade, direitos dos titulares, retencao, incidentes, suboperadores e transferencias internacionais.",
        "Metodo de assinatura e pacote probatorio: versao/hash, identidade, data/hora/fuso, IP, autenticacao, eventos, copia e preservacao do original.",
        "Menores, incapazes, representantes, urgencias, segunda opiniao, retirada de consentimento, recusa, abandono e entrega de prontuario.",
        "Validade e linguagem dos modulos de risco por especialidade, inclusive quando exigir termo proprio, como sedacao.",
        "Compatibilidade entre o texto e a operacao real da Odonto.IA: fornecedores, paises, retencao, backup, RTO/RPO, suporte e uso de IA.",
    ]
    for item in priorities:
        base.add_list_item(doc, item)

    base.add_heading(doc, "4. Referencias oficiais principais", 1)
    sources = [
        "Lei 15.378/2026 - Estatuto dos Direitos do Paciente: https://planalto.gov.br/ccivil_03/_ato2023-2026/2026/lei/l15378.htm",
        "Manual do Prontuario do Paciente em Odontologia - CFO, 2026: https://website.cfo.org.br/wp-content/uploads/2026/03/CFO_Manual_do_Prontuario_Ebook_v2.pdf",
        "Codigo de Etica Odontologica - Resolucao CFO 118/2012: https://website.cfo.org.br/wp-content/uploads/2018/03/codigo_etica.pdf",
        "Lei Geral de Protecao de Dados - Lei 13.709/2018: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709compilado.htm",
        "Guarda e digitalizacao de prontuario - Lei 13.787/2018: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13787.htm",
        "Documentos e assinaturas eletronicas - MP 2.200-2/2001: https://www.planalto.gov.br/ccivil_03/mpv/antigas_2001/2200-2.htm",
        "Marco Civil da Internet - Lei 12.965/2014: https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2014/lei/l12965.htm",
        "Codigo de Defesa do Consumidor - Lei 8.078/1990: https://www.planalto.gov.br/ccivil_03/leis/l8078compilado.htm",
        "ANPD - Resolucao 15/2024, comunicacao de incidentes: https://www.gov.br/anpd/pt-br/canais_atendimento/agente-de-tratamento/comunicado-de-incidente-de-seguranca-cis",
        "ANPD - Resolucao 19/2024, transferencia internacional: https://www.gov.br/anpd/pt-br/acesso-a-informacao/institucional/atos-normativos/regulamentacoes_anpd/resolucao-cd-anpd-no-19-de-23-de-agosto-de-2024",
    ]
    for source in sources:
        base.add_list_item(doc, source)

    path = OUTPUT_DIR / "00-guia-de-envio-ao-advogado.docx"
    finish_document(doc, "Guia de envio ao advogado", path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    main_outputs = build_main_documents()
    module_outputs = build_module_documents()
    all_documents = sorted(main_outputs + module_outputs, key=lambda path: path.name)
    guide = build_guide(all_documents)
    outputs = [guide] + all_documents
    if len(outputs) != 19:
        raise RuntimeError(f"Quantidade inesperada de documentos: {len(outputs)}")
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
