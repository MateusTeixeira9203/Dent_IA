from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "juridico" / "2026-09-03-dossie-juridico-minutas-odonto-ia.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 102, 115)
RED = RGBColor(155, 28, 28)
GOLD = RGBColor(122, 90, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF4CE"
BORDER = "C8D1DC"


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def split_placeholder_text(text: str):
    return re.split(r"(\[\[[^\]]+\]\])", text)


def add_runs(paragraph, text: str, *, bold=False, italic=False, color=None, size=11):
    for piece in split_placeholder_text(text):
        if not piece:
            continue
        run = paragraph.add_run(piece)
        is_placeholder = piece.startswith("[[") and piece.endswith("]]" )
        set_run_font(
            run,
            size=size,
            color=RED if is_placeholder else color,
            bold=True if is_placeholder else bold,
            italic=italic,
        )
        if is_placeholder:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_GOLD)
            run._element.get_or_add_rPr().append(shd)
    return paragraph


def add_para(doc, text="", *, bold=False, italic=False, color=None, size=11,
             before=0, after=6, line=1.25, align=WD_ALIGN_PARAGRAPH.LEFT,
             keep=False, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    add_runs(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_runs(p, text, bold=True, color=BLUE if level < 3 else DARK_BLUE,
             size={1: 16, 2: 13, 3: 12}[level])
    return p


def add_clause(doc, number: str, title: str, paragraphs: list[str], bullets: list[str] | None = None):
    add_heading(doc, f"{number}. {title}", 2)
    for text in paragraphs:
        add_para(doc, text)
    if bullets:
        for item in bullets:
            add_list_item(doc, item)


def add_list_item(doc, text: str, *, numbered=False, level=0):
    style = "Legal Number" if numbered else "Legal Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_runs(p, text)
    return p


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs(p, header, bold=True, color=INK, size=9.5)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for idx, value in enumerate(row_values):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_runs(p, value, size=9.5)
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    add_para(doc, "", after=2, size=2)
    return table


def add_callout(doc, label: str, text: str, *, fill=LIGHT_GRAY, color=INK):
    table = doc.add_table(rows=1, cols=1)
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    add_runs(p, f"{label}: ", bold=True, color=color)
    add_runs(p, text, color=color)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=BORDER)
    add_para(doc, "", after=2, size=2)


def add_signature_block(doc, parties: list[str], include_witnesses=False):
    add_para(doc, "Local: [[CIDADE/UF]]    Data: [[DD/MM/AAAA]]    Hora: [[HH:MM / FUSO]]", before=8)
    rows = []
    for party in parties:
        rows.append(["Assinatura", party])
    if include_witnesses:
        rows.extend([
            ["Testemunha 1", "Nome: [[NOME]] | CPF: [[CPF]] | Assinatura: [[ASSINATURA]]"],
            ["Testemunha 2", "Nome: [[NOME]] | CPF: [[CPF]] | Assinatura: [[ASSINATURA]]"],
        ])
    add_table(doc, ["Campo", "Identificacao"], rows, [1700, 7660])


def add_instrument_header(doc, number: int, title: str, audience: str):
    p = add_para(doc, f"MINUTA {number}", bold=True, color=GOLD, size=10, after=4, keep=True)
    p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    add_para(doc, title.upper(), bold=True, color=INK, size=20, after=4, keep=True)
    add_para(doc, audience, italic=True, color=MUTED, size=10.5, after=12)
    add_callout(
        doc,
        "STATUS",
        "Rascunho para revisao por advogado brasileiro e validacao operacional. Nao publicar nem coletar aceite antes do preenchimento de todos os campos e da aprovacao juridica.",
        fill=LIGHT_GOLD,
        color=RED,
    )


def patch_numbering(doc):
    numbering = doc.part.numbering_part.element

    def make_abstract(num_id: int, bullet: bool):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(num_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if bullet:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Symbol")
            fonts.set(qn("w:hAnsi"), "Symbol")
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(num_id))
        num.append(abstract_ref)
        numbering.append(num)

    make_abstract(81, True)
    make_abstract(82, False)

    styles = doc.styles
    for name, num_id in (("Legal Bullet", 81), ("Legal Number", 82)):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        p_pr = style.element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    configs = {
        "Heading 1": (16, BLUE, 14, 8),
        "Heading 2": (13, BLUE, 11, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in configs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    patch_numbering(doc)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def setup_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_runs(p, "ODONTO.IA  |  DOSSIÊ JURIDICO PRELIMINAR  |  CONFIDENCIAL", bold=True, color=MUTED, size=8.5)
        footer = section.footer
        fp = footer.paragraphs[0]
        add_page_field(fp)


def build_front_matter(doc):
    add_para(doc, "ODONTO.IA", bold=True, color=GOLD, size=10, after=28)
    add_para(doc, "Dossiê jurídico preliminar", bold=True, color=INK, size=28, after=4)
    add_para(doc, "Ecossistema de contratos, consentimentos, privacidade e prova eletrônica", color=DARK_BLUE, size=15, after=18)
    add_para(doc, "Versao 0.1 | 3 de setembro de 2026", bold=True, color=MUTED, size=10.5, after=28)
    add_callout(
        doc,
        "NAO PUBLICAR",
        "Material preparatorio para revisao de advogado, validacao clinica por especialistas e confirmacao tecnica da Odonto.IA. As minutas contem campos [[PREENCHER]] e decisoes em aberto.",
        fill=LIGHT_GOLD,
        color=RED,
    )
    add_para(doc, "Objetivo", bold=True, color=BLUE, size=13, before=10, after=4)
    add_para(doc, "Entregar uma base coerente para proteger a Odonto.IA dentro dos limites legais e, ao mesmo tempo, aumentar a qualidade da documentacao do cirurgiao-dentista. O dossie nao promete imunidade: organiza deveres, informacao, consentimento, rastreabilidade e resposta a incidentes.")
    add_para(doc, "Premissa central", bold=True, color=BLUE, size=13, before=8, after=4)
    add_para(doc, "O melhor mecanismo de defesa nao e uma clausula ampla de isencao. E a capacidade de demonstrar o que foi explicado, qual alternativa foi escolhida, quais riscos eram relevantes para aquele caso, quem assinou, quando assinou e se o documento permaneceu integro.")

    doc.add_page_break()
    add_heading(doc, "1. Recomendacao executiva", 1)
    add_para(doc, "A Odonto.IA deve operar com duas camadas juridicas separadas:")
    add_list_item(doc, "Camada plataforma: contratacao SaaS, privacidade de usuarios, tratamento de dados de pacientes, seguranca, subprocessadores e limites da inteligencia artificial.")
    add_list_item(doc, "Camada clinica: contrato de prestacao odontologica, plano/orcamento, TCLE individualizado, registros de evolucao e intercorrencias, orientacoes, conclusao, recusa e entrega de documentos.")
    add_para(doc, "Misturar as duas camadas enfraquece a prova. O paciente nao contrata a Odonto.IA; a clinica nao deve tentar obter, no cadastro do dentista, consentimento em nome de pacientes que ainda nem foram identificados.")

    add_heading(doc, "2. Arquitetura documental proposta", 1)
    add_table(
        doc,
        ["Instrumento", "Momento", "Quem aceita/assina", "Funcao"],
        [
            ["Termos de Uso e Contrato SaaS", "Primeiro acesso e nova versao material", "Contratante ou representante autorizado; cada usuario adere ao uso aceitavel", "Reger a relacao Odonto.IA-cliente"],
            ["Aviso de Privacidade de Usuarios", "Cadastro e acesso permanente", "Ciencia do usuario; consentimento apenas onde for realmente a base legal", "Transparencia sobre dados de dentistas/equipe"],
            ["Acordo de Tratamento de Dados (DPA)", "Contratacao da clinica", "Titular/administrador com poderes", "Controlador-operador, seguranca, incidentes, suboperadores"],
            ["Aviso de Privacidade do Paciente", "Cadastro ou primeiro atendimento", "Disponibilizado pela clinica ao paciente", "Explicar o tratamento de dados clinicos"],
            ["Contrato de Servicos Odontologicos", "Inicio do tratamento", "Clinica/profissional + paciente/responsavel", "Objeto, deveres, prazo, valores, rescisao"],
            ["Plano e Aceite de Orcamento", "Antes da execucao dos itens", "Paciente/responsavel + profissional", "Congelar itens, alternativas, valores e condicoes"],
            ["TCLE nucleo + modulo de risco", "Antes do procedimento", "Paciente/responsavel + dentista; testemunha conforme risco", "Consentimento clinico individualizado"],
            ["Orientacoes e Conclusao", "Ao final da sessao/procedimento", "Paciente/responsavel + executor", "Registrar o realizado, intercorrencias e cuidados"],
            ["Recusa/Interrupcao/Abandono", "Quando houver recusa ou quebra de continuidade", "Paciente/responsavel + profissional", "Registrar recomendacao, riscos e continuidade segura"],
            ["Audio, voz, imagem e transcricao", "Antes de capturar paciente identificavel", "Paciente/responsavel, de forma granular", "Separar documentacao clinica de marketing/pesquisa"],
        ],
        [2400, 1900, 2300, 2760],
    )

    add_heading(doc, "3. Cadastro e aceite no primeiro acesso", 1)
    add_callout(doc, "CORRECAO", "Para cirurgiao-dentista, o registro profissional e CRO-UF e numero. CRM so deve existir se o produto admitir medicos em fluxo proprio.", fill=LIGHT_BLUE)
    add_para(doc, "Dados recomendados, observada a minimizacao da LGPD:")
    for item in [
        "Nome civil completo e nome social, quando aplicavel; e-mail verificado; telefone para seguranca/recuperacao.",
        "CRO-UF e numero; especialidades registradas, se usadas no produto; declaracao de regularidade e autorizacao para verificacao em fonte publica.",
        "Se pessoa juridica: razao social, CNPJ, numero EPAO/CRO, endereco, responsavel tecnico e prova/declaracao de poderes do administrador.",
        "Papel na clinica: titular, administrador, cirurgiao-dentista, secretaria ou outro; permissoes nao podem nascer apenas do que o proprio usuario declara.",
        "CPF apenas quando necessario a identificacao contratual, fiscal, antifraude ou assinatura; a finalidade deve ser documentada.",
    ]:
        add_list_item(doc, item)
    add_para(doc, "Aceites separados no cadastro:")
    add_list_item(doc, "Obrigatorio: 'Li e aceito os Termos de Uso e o Contrato SaaS, versao X'.")
    add_list_item(doc, "Obrigatorio como ciencia, nao como consentimento generico: 'Li o Aviso de Privacidade'.")
    add_list_item(doc, "Somente titular/administrador: 'Possuo poderes para contratar e aprovo o Acordo de Tratamento de Dados em nome da clinica'.")
    add_list_item(doc, "Opcional e desmarcado: marketing, pesquisa ou usos secundarios que realmente dependam de consentimento.")
    add_para(doc, "O acesso de um membro convidado nao deve permitir que ele aceite o DPA em nome da clinica. Ele adere aos deveres de confidencialidade, seguranca e uso aceitavel.")

    add_heading(doc, "4. Pacote minimo de prova do aceite eletronico", 1)
    for item in [
        "Identificador do documento, tipo, versao, hash SHA-256 do conteudo final e PDF congelado.",
        "Identificacao do signatario, qualidade em que assina e, se representante, vinculo declarado e documento.",
        "Data, hora, fuso, IP, user agent, sessao autenticada e metodo de verificacao de identidade.",
        "Eventos da jornada: abriu, rolou/visualizou, confirmou declaracoes, assinou, recebeu copia e eventual revogacao/recusa.",
        "Vinculos server-side com clinica, paciente, dentista, ficha, procedimento e orcamento; nunca confiar nos IDs ou valores enviados pelo navegador.",
        "Assinaturas do paciente e do profissional. Para contratos de cobranca, avaliar duas testemunhas ou provedor de assinatura com integridade juridicamente adequada.",
        "Trilha de auditoria imutavel, controle de versao, vedacao de UPDATE/DELETE pelo usuario comum e processo de retificacao por adendo, sem apagar o original.",
        "Entrega de uma copia ao signatario e registro dessa entrega.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "5. O que os contratos nao podem prometer", 1)
    for item in [
        "Que o TCLE elimina responsabilidade por erro, culpa, defeito do servico, falha de seguranca ou descumprimento do dever de informar.",
        "Que a inteligencia artificial diagnostica, decide ou substitui a revisao do profissional.",
        "Que todo tratamento de dado de saude depende de consentimento. A base legal deve ser escolhida por finalidade.",
        "Que a assinatura desenhada e ICP-Brasil ou assinatura qualificada, se nao for.",
        "Que dados serao apagados imediatamente quando houver obrigacao legal, regulatoria ou probatoria de guarda.",
        "SLA, backup, localizacao de dados, RTO/RPO ou prazo de exportacao que a operacao tecnica ainda nao consiga cumprir.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "6. Decisoes que o advogado e a equipe tecnica devem fechar", 1)
    add_table(
        doc,
        ["Tema", "Decisao pendente", "Responsavel"],
        [
            ["Parte contratante", "Pessoa fisica, clinica PJ ou ambas; capacidade de quem aceita", "Juridico + comercial"],
            ["Consumerista", "Quando o CDC pode incidir sobre cliente profissional/PJ e como redigir limites", "Juridico"],
            ["Retencao/exportacao", "Prazos reais apos cancelamento e fluxo de portabilidade", "Tecnico + juridico"],
            ["Suboperadores", "Entidade juridica, pais, regiao, dado enviado, retencao e mecanismo internacional", "Tecnico + privacidade"],
            ["IA", "Garantir contratualmente ausencia de treinamento com dados clinicos e politicas de retencao", "Tecnico + compras"],
            ["Audio ambiente", "Bloquear captura do paciente por padrao ou coletar autorizacao especifica", "Produto + juridico"],
            ["Assinatura", "Pad desenhado, OTP, assinatura avancada ou ICP-Brasil por classe de risco", "Juridico + produto"],
            ["Riscos clinicos", "Validar modulos por especialistas e definir quando novo TCLE e obrigatorio", "Diretor clinico"],
            ["Incidentes", "SLA interno da Odonto.IA para avisar a clinica e canal 24x7", "Seguranca + juridico"],
        ],
        [1900, 5560, 1900],
    )


def build_terms(doc):
    add_instrument_header(doc, 1, "Termos de Uso e Contrato de Licenca SaaS", "Relacao entre a Odonto.IA e o profissional ou clinica contratante")
    add_para(doc, "CONTRATADA: [[ODONTOIA_RAZAO_SOCIAL]], pessoa juridica inscrita no CNPJ sob nº [[CNPJ]], com sede em [[ENDERECO]], contato [[SUPORTE]].")
    add_para(doc, "CONTRATANTE: pessoa fisica ou juridica identificada no cadastro e na Ordem de Contratacao, representada por usuario que declara possuir poderes para contratar.")
    add_para(doc, "Ao selecionar 'Aceitar e continuar', o Contratante celebra este instrumento e cada Usuario Autorizado adere aos deveres de uso, confidencialidade e seguranca aplicaveis ao seu perfil.")
    add_clause(doc, "1", "Definicoes", ["Para este instrumento: Plataforma e o software Odonto.IA; Contratante e o cliente da assinatura; Clinica e o estabelecimento controlador dos dados de pacientes; Usuario Autorizado e a pessoa convidada; Dados Clinicos sao informacoes de saude e prontuario; Recursos de IA sao funcoes assistivas de transcricao, organizacao, estruturacao ou sugestao."])
    add_clause(doc, "2", "Objeto e licenca", ["A CONTRATADA concede, durante a vigencia do plano, licenca limitada, revogavel, nao exclusiva, intransferivel e sem sublicenciamento para uso interno profissional da Plataforma, conforme funcionalidades, limites e preco da Ordem de Contratacao vigente.", "A licenca nao transfere codigo-fonte, modelos, marcas, documentacao proprietaria ou outros direitos de propriedade intelectual."])
    add_clause(doc, "3", "Natureza do servico e da inteligencia artificial", ["A Plataforma e ferramenta de apoio administrativo e documental. Nao presta atendimento odontologico, nao estabelece diagnostico, prognostico ou plano terapeutico e nao substitui o julgamento do cirurgiao-dentista.", "Saidas de IA podem conter omissoes ou erros. O profissional deve revisar, corrigir, confirmar e assumir a autoria clinica antes de salvar, assinar, emitir ou compartilhar qualquer documento. A ausencia de revisao constitui uso indevido.", "A Plataforma nao deve ser utilizada como sistema de emergencia nem como unica fonte de informacao indispensavel a conduta imediata. A Clinica manterá procedimento de contingencia compativel com sua atividade."])
    add_clause(doc, "4", "Elegibilidade e qualificacao profissional", ["O usuario declara ter capacidade civil e, quando atuar profissionalmente, possuir registro compativel e regular no conselho competente. O cirurgiao-dentista informara CRO-UF e numero verdadeiros e manterá os dados atualizados.", "Quem aceitar em nome de pessoa juridica declara possuir poderes suficientes. Informacao falsa ou uso por pessoa nao habilitada autoriza suspensao cautelar, preservados dados, direitos dos pacientes e deveres legais."])
    add_clause(doc, "5", "Obrigacoes do Contratante e dos Usuarios", ["O Contratante e responsavel por configurar perfis, revisar acessos, retirar usuarios desligados, manter dados de cadastro atualizados e assegurar que cada pessoa utilize credencial individual."], [
        "Inserir somente dados licitamente obtidos e necessarios a finalidade assistencial ou administrativa declarada.",
        "Nao compartilhar senhas, contornar controles de acesso, testar vulnerabilidades sem autorizacao ou acessar outra clinica.",
        "Nao usar a Plataforma para exercicio profissional sem habilitacao, fraude, discriminacao, publicidade irregular ou conteudo ilicito.",
        "Conferir documentos clinicos, assinaturas, orcamentos, prescricoes e comunicacoes antes do uso externo.",
        "Comunicar imediatamente suspeita de acesso indevido, perda de credencial ou incidente.",
    ])
    add_clause(doc, "6", "Dados de pacientes e LGPD", ["No tratamento realizado em nome da Clinica, a Clinica atua como Controladora e a CONTRATADA como Operadora, conforme o Acordo de Tratamento de Dados anexo. A qualificacao podera variar quando a CONTRATADA tratar dados para obrigacoes proprias, seguranca, faturamento ou defesa de direitos, devendo isso constar do Aviso de Privacidade.", "O Contratante determina finalidades e bases legais, presta informacoes aos pacientes e responde pela autorizacao de acesso por sua equipe. A CONTRATADA trata dados segundo instrucoes documentadas, salvo obrigacao legal em contrario."])
    add_clause(doc, "7", "Confidencialidade e sigilo", ["As partes manterao confidenciais dados clinicos, credenciais, documentos, informacoes comerciais e tecnicas. O dever permanece apos o termino enquanto a informacao estiver protegida por lei, sigilo profissional ou sua natureza.", "Acesso interno sera limitado a pessoas que necessitem da informacao, vinculadas a deveres de confidencialidade e treinadas para sua funcao."])
    add_clause(doc, "8", "Suboperadores e integracoes", ["A CONTRATADA podera utilizar suboperadores listados no Anexo do DPA para hospedagem, banco, autenticacao, inteligencia artificial, comunicacao, cobranca e seguranca. Mudanca material seguirá o mecanismo de aviso e oposicao definido no DPA.", "Integracoes opcionais ativadas pelo Contratante podem estar sujeitas a termos proprios. A ativacao nao autoriza uso alem das finalidades informadas."])
    add_clause(doc, "9", "Planos, cobranca e tributos", ["Preco, ciclo, limites, teste, renovacao e meios de pagamento constam da Ordem de Contratacao ou checkout. Alteracoes de preco serao informadas com antecedencia de [[PRAZO]].", "A inadimplencia pode gerar restricao de funcionalidades nao essenciais e, apos aviso, suspensao. O tratamento de prontuario, exportacao e continuidade assistencial observara os limites legais e a politica de saida; dados nao serao retidos como meio coercitivo de cobranca."])
    add_clause(doc, "10", "Disponibilidade, manutencao e suporte", ["A CONTRATADA empregara esforcos tecnicos compativeis com o plano e o nivel de risco. Janelas, suporte, backup, RTO, RPO e exclusoes de SLA devem constar de [[ANEXO_TECNICO/SLA]] e somente podem ser prometidos apos validacao operacional.", "Podera haver manutencao programada, atualizacoes de seguranca e indisponibilidade por terceiros ou forca maior. A CONTRATADA informara incidentes relevantes e adotara medidas de mitigacao conforme o DPA."])
    add_clause(doc, "11", "Propriedade e uso de dados", ["O Contratante e os titulares mantem os direitos sobre seus dados. A CONTRATADA nao adquire propriedade sobre prontuarios ou documentos clinicos.", "Dados Clinicos identificados nao serao usados para treinamento de modelos de proposito geral ou finalidade propria incompatível sem instrumento especifico, base legal valida e autorizacoes exigiveis. Metricas agregadas somente poderao ser utilizadas quando nao permitirem identificacao ou reidentificacao razoavel."])
    add_clause(doc, "12", "Suspensao e seguranca", ["Acesso pode ser suspenso cautelarmente quando houver risco concreto a seguranca, uso ilicito, violacao grave, ordem de autoridade ou inadimplencia nos termos da clausula 9. Sempre que viavel, havera aviso, escopo proporcional e meio de contestacao.", "A suspensao preservara evidencias, deveres de sigilo, disponibilidade de dados conforme a lei e medidas necessarias a continuidade segura do atendimento."])
    add_clause(doc, "13", "Vigencia, cancelamento e saida", ["Este instrumento vigorara enquanto houver assinatura ativa ou obrigacoes remanescentes. O cancelamento seguira a Ordem de Contratacao, sem prejuizo do pagamento devido e do direito de encerrar renovacao.", "Por [[PRAZO_DE_EXPORTACAO]] apos o termino, o Contratante podera solicitar exportacao em formato [[FORMATOS]]. Depois, a CONTRATADA eliminara ou anonimizará dados conforme instrucoes e lei, ressalvadas copias tecnicas temporarias e guarda obrigatoria, tudo conforme o DPA."])
    add_clause(doc, "14", "Responsabilidade", ["Cada parte responde pelos danos que causar por violacao de lei, contrato ou dever de seguranca. Nenhuma disposicao exclui responsabilidade que nao possa ser afastada, inclusive por dolo, culpa grave, violacao de confidencialidade, infracao de protecao de dados ou direitos do consumidor quando aplicaveis.", "Ressalvadas as hipoteses nao limitaveis, eventual limite B2B de responsabilidade devera ser definido pelo advogado em [[LIMITE_E_EXCECOES]], de forma proporcional ao risco e ao valor contratado. Nao ha transferencia ao profissional de falhas proprias da Plataforma, nem a CONTRATADA assume decisoes clinicas do usuario."])
    add_clause(doc, "15", "Indenizacao entre as partes", ["A parte que, por ato proprio comprovado, causar reclamacao de terceiro indenizara a outra pelos prejuizos diretos e razoaveis, assegurados aviso, cooperacao, direito de defesa e vedacao a acordo que imponha obrigacao sem anuencia. A clausula nao cria renuncia de direitos de pacientes ou titulares."])
    add_clause(doc, "16", "Alteracoes dos termos", ["Mudancas meramente formais podem ser publicadas com aviso. Mudanca material de finalidade, responsabilidade, preco, dados ou direitos exige comunicacao destacada e novo aceite quando juridicamente necessario. Cada versao tera data, hash e historico."])
    add_clause(doc, "17", "Lei, solucao de conflitos e foro", ["Aplica-se a lei brasileira. As partes buscarao solucao pelo suporte e, se desejarem, mediacao. O foro sera [[FORO_VALIDADO]], sem afastar foro legalmente competente nem direitos consumeristas eventualmente aplicaveis."])
    add_clause(doc, "18", "Disposicoes finais", ["A tolerancia nao implica renuncia. Nulidade parcial nao invalida as demais clausulas quando puderem subsistir. Ordem de Contratacao, DPA, Aviso de Privacidade, Anexo de Seguranca e lista de Suboperadores integram o contrato; em conflito sobre dados pessoais, prevalece a regra mais protetiva e especifica, sujeita a revisao juridica."])
    add_callout(doc, "DECLARACAO DE ACEITE", "Declaro que li este instrumento, possuo poderes quando atuo por pessoa juridica, informei dados verdadeiros e aceito a versao [[VERSAO/HASH]]. Reconheco que recursos de IA sao assistivos e que documentos clinicos exigem revisao e validacao profissional antes do uso.", fill=LIGHT_BLUE)


def build_user_privacy(doc):
    add_instrument_header(doc, 2, "Aviso de Privacidade de Usuarios da Plataforma", "Dentistas, administradores, secretarias, convidados e contatos comerciais")
    add_clause(doc, "1", "Quem trata os dados", ["[[ODONTOIA_RAZAO_SOCIAL]], CNPJ [[CNPJ]], e a controladora dos dados de cadastro, seguranca, faturamento, suporte e relacionamento descritos neste Aviso. Contato do encarregado/canal de privacidade: [[CANAL_PRIVACIDADE]].", "Para dados de pacientes tratados sob instrucao da Clinica, a Odonto.IA atua predominantemente como operadora; esse fluxo e regido pelo DPA e pelo aviso da propria Clinica."])
    add_clause(doc, "2", "Dados tratados", ["Podemos tratar identificacao e contato; CRO-UF e informacoes profissionais; clinicas e papeis; credenciais e registros de autenticacao; IP, dispositivo e registros de acesso; eventos de seguranca e auditoria; plano, cobranca e dados fiscais; suporte e comunicacoes; preferencias; e dados necessarios a integracoes ativadas."])
    add_clause(doc, "3", "Finalidades e bases legais", ["As bases serao definidas por finalidade, nao por um consentimento amplo."], [
        "Executar o contrato e procedimentos preliminares: criar conta, autenticar, disponibilizar funcoes, cobrar e prestar suporte.",
        "Cumprir obrigacoes legais ou regulatorias: documentos fiscais, registros obrigatorios e resposta a autoridades competentes.",
        "Interesse legitimo, apos avaliacao: seguranca, prevencao a fraude, melhoria operacional e comunicacao com clientes, com direito de oposicao quando aplicavel.",
        "Exercicio regular de direitos: preservar evidencias e atuar em processos.",
        "Consentimento: somente para finalidades opcionais claramente apresentadas, como determinadas comunicacoes de marketing, podendo ser revogado.",
    ])
    add_clause(doc, "4", "Compartilhamento", ["Dados podem ser compartilhados com fornecedores de infraestrutura, autenticacao, pagamentos, comunicacao, suporte e seguranca; com a Clinica a que o usuario pertence; e com autoridades quando houver obrigacao ou ordem valida. A lista material de fornecedores e suas finalidades ficara em [[URL_SUBOPERADORES]]."])
    add_clause(doc, "5", "Transferencias internacionais", ["Alguns fornecedores podem tratar dados fora do Brasil. A Odonto.IA informara paises, finalidades, responsabilidades e mecanismos de transferencia em [[URL_TRANSFERENCIAS]]. Quando forem usadas clausulas-padrao da ANPD, seu texto integral sera incorporado sem alteracao aos instrumentos aplicaveis."])
    add_clause(doc, "6", "Retencao", ["Os prazos consideram contrato, obrigacoes legais, prevencao a fraude e defesa de direitos. Registros de acesso a aplicacao serao mantidos pelo prazo legal aplicavel. A tabela completa de retencao deve ser publicada em [[TABELA_RETENCAO]]. Dados desnecessarios serao eliminados ou anonimizados com seguranca."])
    add_clause(doc, "7", "Direitos", ["O titular pode solicitar confirmacao, acesso, correcao, informacao sobre compartilhamento, portabilidade quando aplicavel, revisao de decisoes automatizadas, oposicao, revogacao de consentimento e eliminacao nas hipoteses legais. Pedidos serao recebidos em [[CANAL_PRIVACIDADE]] e podem exigir verificacao proporcional de identidade."])
    add_clause(doc, "8", "Seguranca e incidentes", ["Adotamos medidas tecnicas e administrativas proporcionais ao risco, incluindo controle de acesso, isolamento entre clinicas, registros, criptografia conforme arquitetura, backup e resposta a incidentes. Nenhum sistema e isento de risco. Incidentes relevantes serao tratados e comunicados conforme a legislacao e os papeis de controlador/operador."])
    add_clause(doc, "9", "Cookies e tecnologias semelhantes", ["Cookies estritamente necessarios podem ser usados para sessao e seguranca. Analiticos, publicidade ou tecnologias opcionais devem constar de [[POLITICA_COOKIES]] e respeitar escolhas do usuario."])
    add_clause(doc, "10", "Atualizacoes", ["Mudancas materiais serao destacadas. A versao, data e historico ficarao disponiveis. Se a finalidade depender de consentimento, mudanca relevante exigira nova manifestacao quando cabivel."])


def build_dpa(doc):
    add_instrument_header(doc, 3, "Acordo de Tratamento de Dados Pessoais (DPA)", "Anexo ao Contrato SaaS entre a Clinica controladora e a Odonto.IA operadora")
    add_para(doc, "CONTROLADORA: [[CLINICA_RAZAO_SOCIAL/PROFISSIONAL]], CNPJ/CPF [[NUMERO]], contato de privacidade [[CONTATO]].")
    add_para(doc, "OPERADORA: [[ODONTOIA_RAZAO_SOCIAL]], CNPJ [[CNPJ]], contato [[CANAL_PRIVACIDADE]].")
    add_clause(doc, "1", "Objeto, papeis e instrucoes", ["A OPERADORA tratara dados pessoais descritos no Anexo A exclusivamente para disponibilizar, proteger e manter a Plataforma e as funcionalidades ativadas, segundo instrucoes documentadas da CONTROLADORA e a legislacao.", "A CONTROLADORA define finalidades e bases legais da assistencia, informa os titulares e autoriza usuarios. A OPERADORA informara se considerar uma instrucao manifestamente ilicita e podera suspender somente a operacao afetada enquanto aguarda ajuste."])
    add_clause(doc, "2", "Pessoal autorizado e confidencialidade", ["A OPERADORA limitara o acesso a pessoas que necessitem dele, sujeitas a confidencialidade e treinamento. A CONTROLADORA aplicara o mesmo principio a sua equipe e revisara acessos periodicamente."])
    add_clause(doc, "3", "Seguranca", ["A OPERADORA mantera medidas proporcionais a dados de saude e ao risco, detalhadas no Anexo B, incluindo segregacao multi-clinica, menor privilegio, autenticacao, criptografia em transito e conforme arquitetura em repouso, logs, gestao de vulnerabilidades, backup, continuidade e resposta a incidentes.", "O Anexo B deve descrever apenas controles efetivamente implementados e indicar responsavel, evidencia, frequencia e excecoes."])
    add_clause(doc, "4", "Suboperadores", ["A CONTROLADORA concede autorizacao geral aos suboperadores do Anexo C, condicionada a contratos com protecao equivalente. A OPERADORA avisara alteracao material com [[PRAZO]] de antecedencia. A CONTROLADORA podera apresentar objecao fundamentada em protecao de dados; as partes buscarao alternativa viavel ou encerramento da funcao afetada."])
    add_clause(doc, "5", "Direitos dos titulares", ["A OPERADORA encaminhara pedidos recebidos diretamente e prestara assistencia razoavel para acesso, copia, correcao, restricao, portabilidade, informacao e eliminacao, respeitados sigilo, identidade e obrigacoes de guarda. Prazo operacional interno: [[SLA_TITULARES]]."])
    add_clause(doc, "6", "Incidentes", ["A OPERADORA comunicara a CONTROLADORA sem demora injustificada apos confirmar incidente envolvendo seus dados, pelo canal [[CANAL_24X7]], com meta contratual de [[EX.: 24 HORAS]], a ser validada tecnicamente. Fornecera natureza, titulares/dados afetados, medidas, riscos, ponto de contato e atualizacoes.", "A CONTROLADORA decide e realiza comunicacoes a ANPD e titulares quando obrigatorias. A OPERADORA preservara evidencias e cooperara. Nenhuma parte fara anuncio em nome da outra sem base legal ou alinhamento, salvo obrigacao independente."])
    add_clause(doc, "7", "Avaliacoes e auditoria", ["A OPERADORA fornecera informacoes razoaveis de conformidade, relatorios ou certificacoes disponiveis. Auditoria adicional dependera de justificativa, confidencialidade, escopo, antecedencia e ausencia de risco a outros clientes. Achado critico comprovado sera corrigido segundo plano acordado."])
    add_clause(doc, "8", "Transferencia internacional", ["Transferencias observarao finalidade, necessidade, base legal e mecanismo valido. Se baseadas em clausulas-padrao contratuais da ANPD, o texto oficial sera anexado integralmente e sem alteracao, com preenchimento das secoes permitidas. Paises, importadores, categorias e salvaguardas constarao do Anexo C e do aviso publico."])
    add_clause(doc, "9", "Termino, devolucao e eliminacao", ["Encerrado o servico, a OPERADORA disponibilizara exportacao por [[PRAZO]] e, depois, eliminara ou anonimizara dados conforme instrucao, salvo guarda legal, defesa de direitos e ciclos tecnicos de backup descritos em [[PRAZO_BACKUP]]. Dados preservados ficarao isolados e sem uso operacional."])
    add_clause(doc, "10", "Responsabilidade e precedencia", ["Cada parte responde por seu papel e por instrucoes/operacoes sob seu controle. Este DPA nao reduz direitos de titulares. Em conflito com o Contrato SaaS sobre dados pessoais, prevalece este DPA; clausulas-padrao da ANPD prevalecem no escopo da transferencia internacional."])
    add_heading(doc, "Anexo A - Descricao do tratamento", 2)
    add_table(doc, ["Campo", "Descricao"], [
        ["Titulares", "Pacientes, responsaveis/representantes, dentistas, equipe, contatos e pagadores quando aplicavel"],
        ["Dados", "Identificacao, contato, agenda, saude, anamnese, prontuario, imagens/documentos, voz/audio quando ativado, financeiro do tratamento, assinaturas e auditoria"],
        ["Operacoes", "Coleta, recepcao, organizacao, transcricao, estruturacao, armazenamento, consulta, geracao documental, transmissao autorizada, backup, exportacao e eliminacao"],
        ["Finalidades", "Gestao clinica e administrativa, documentacao, comunicacao ativada, seguranca, suporte e cumprimento das instrucoes da Clinica"],
        ["Duracao", "Vigencia contratual mais prazos de exportacao, backup e guarda definidos neste DPA"],
        ["Dados especiais", "Dados de saude; voz/imagem; dados de menores e incapazes; assinatura; eventualmente dados financeiros"],
    ], [1900, 7460])
    add_heading(doc, "Anexo B - Medidas de seguranca a validar", 2)
    for item in [
        "RLS e isolamento por clinica; testes com duas contas e monitoramento de falhas de autorizacao.",
        "MFA para perfis privilegiados, gestao de sessao, rotacao de segredos e revogacao de usuarios.",
        "Criptografia em transito e em repouso, com inventario de chaves e responsabilidades.",
        "Logs sem conteudo clinico desnecessario; trilha de auditoria para leitura, alteracao, exportacao e assinatura.",
        "Backups testados, RTO/RPO documentados, restauracao e plano de continuidade.",
        "Desenvolvimento seguro, gestao de vulnerabilidades, dependencias, segregacao de ambientes e resposta a incidentes.",
        "Politica de suporte que proiba copiar prontuario para canais nao aprovados.",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "Anexo C - Suboperadores preliminares", 2)
    add_callout(doc, "VALIDACAO OBRIGATORIA", "A tabela abaixo reflete integracoes encontradas no codigo, nao confirma entidade contratual, regiao, retencao ou mecanismo internacional. Juridico e engenharia devem preencher antes da publicacao.", fill=LIGHT_GOLD, color=RED)
    add_table(doc, ["Fornecedor/servico", "Finalidade", "Dados previstos", "Pais/mecanismo"], [
        ["Supabase", "Banco, autenticacao e armazenamento", "Conta, clinica, prontuario e documentos", "[[VALIDAR ENTIDADE/REGIAO/TRANSFERENCIA]]"],
        ["Vercel", "Hospedagem, funcoes e entrega", "Requisicoes, logs tecnicos e dados processados pelas funcoes", "[[VALIDAR]]"],
        ["Google Gemini", "Estruturacao e extracao por IA", "Relato/conteudo clinico minimizado", "[[VALIDAR RETENCAO/TREINO/TRANSFERENCIA]]"],
        ["Groq", "Transcricao e geracao assistiva", "Audio ou texto estritamente necessario", "[[VALIDAR RETENCAO/TREINO/TRANSFERENCIA]]"],
        ["Upstash", "Rate limit/cache", "Identificadores tecnicos; vedar conteudo clinico", "[[VALIDAR]]"],
        ["Resend", "E-mail transacional", "Conta e mensagens; vedar conteudo clinico por padrao", "[[VALIDAR]]"],
        ["Meta WhatsApp", "Mensagens e envio ativado pela clinica", "Contato e conteudo selecionado", "[[VALIDAR E INFORMAR AO PACIENTE]]"],
        ["Stripe", "Cobranca da assinatura SaaS", "Dados do cliente contratante; sem prontuario", "[[VALIDAR]]"],
        ["Google Calendar", "Integracao opcional de agenda", "Eventos e identificadores configurados", "[[VALIDAR]]"],
    ], [1800, 2100, 2900, 2560])


def build_patient_privacy(doc):
    add_instrument_header(doc, 4, "Aviso de Privacidade do Paciente", "Modelo para personalizacao e publicacao pela Clinica controladora")
    add_callout(doc, "ATENCAO", "Este aviso pertence a Clinica, nao deve apresentar a Odonto.IA como controladora de toda a assistencia. A Clinica precisa adaptar finalidades, bases legais, fornecedores e canais reais.", fill=LIGHT_GOLD, color=RED)
    add_para(doc, "CONTROLADOR: [[CLINICA/PROFISSIONAL]], CNPJ/CPF [[NUMERO]], endereco [[ENDERECO]], responsavel tecnico [[NOME/CRO]], canal de privacidade [[CONTATO]].")
    add_clause(doc, "1", "Por que tratamos seus dados", ["Utilizamos seus dados para identificar voce, avaliar e acompanhar sua saude bucal, elaborar e manter prontuario, planejar e executar tratamentos, emitir documentos, gerir agenda e pagamentos, comunicar orientacoes, cumprir obrigacoes e exercer direitos."])
    add_clause(doc, "2", "Quais dados", ["Podemos tratar identificacao, contato, dados de responsavel, historico e condicoes de saude, medicamentos, alergias, exames, imagens, odontograma, procedimentos, evolucao, intercorrencias, orcamentos, pagamentos, assinaturas e comunicacoes. Dados de saude sao sensiveis e recebem protecao reforcada."])
    add_clause(doc, "3", "Bases legais", ["A Clinica escolhera a base legal adequada a cada finalidade, que pode incluir tutela da saude por profissionais/servicos de saude, cumprimento de obrigacao legal ou regulatoria, execucao de contrato, exercicio regular de direitos, protecao da vida e consentimento especifico quando exigido. Consentimento clinico para um procedimento nao e o mesmo que consentimento LGPD para uso opcional de imagem ou marketing."])
    add_clause(doc, "4", "Uso da Odonto.IA e de inteligencia artificial", ["A Clinica utiliza a Odonto.IA para organizar, armazenar e gerar documentos do atendimento. Recursos de IA podem transcrever ou estruturar informacoes, mas nao substituem o cirurgiao-dentista. O profissional revisa o conteudo antes de integra-lo ao prontuario. Fornecedores e transferencias aplicaveis devem ser listados em [[URL/LISTA]]."])
    add_clause(doc, "5", "Compartilhamento", ["Dados podem ser acessados pela equipe autorizada, profissionais envolvidos no cuidado, laboratorios e servicos auxiliares quando necessarios, operadoras de saude, fornecedores tecnologicos e autoridades nas hipoteses legais. Nao comercializamos dados de saude."])
    add_clause(doc, "6", "Audio, voz e imagem", ["Gravacao de audio/video, transcricao de voz do paciente e uso de imagem para ensino, divulgacao, pesquisa ou marketing serao tratados separadamente, com finalidade clara e opcao real. A recusa a uso opcional nao reduz a qualidade do tratamento. Registros estritamente clinicos seguem as bases e prazos aplicaveis ao prontuario."])
    add_clause(doc, "7", "Guarda e seguranca", ["O prontuario e documentos serao guardados pelo prazo legal e probatorio aplicavel, com controles contra acesso, alteracao e destruicao nao autorizados. A eliminacao depende da natureza do documento e das obrigacoes profissionais; por isso, nem todo pedido de exclusao pode ser atendido integralmente."])
    add_clause(doc, "8", "Seus direitos", ["Voce pode pedir informacoes, acesso e copia gratuita do prontuario, correcao, esclarecimento sobre compartilhamento, oposicao e demais direitos legais. Pode retirar consentimento clinico antes do procedimento e revogar autorizacoes opcionais, sem apagar fatos clinicos ou registros cuja guarda seja obrigatoria. Contato: [[CANAL]]."])
    add_clause(doc, "9", "Menores, incapacidade e acessibilidade", ["Representantes serao identificados e a participacao do paciente sera respeitada conforme sua capacidade. A Clinica adotara meios acessiveis, interprete ou apoio quando necessario para que a informacao seja compreendida."])


def build_service_contract(doc):
    add_instrument_header(doc, 5, "Contrato de Prestacao de Servicos Odontologicos", "Clinica/profissional e paciente ou responsavel legal")
    add_para(doc, "CONTRATADA: [[CLINICA/PROFISSIONAL]], CNPJ/CPF [[NUMERO]], EPAO/CRO [[REGISTRO]], endereco [[ENDERECO]], responsavel tecnico [[NOME/CRO]].")
    add_para(doc, "CONTRATANTE: [[PACIENTE OU RESPONSAVEL]], CPF [[CPF]], endereco [[ENDERECO]], contato [[CONTATO]], na qualidade de [[PACIENTE/RESPONSAVEL]].")
    add_para(doc, "PACIENTE BENEFICIARIO, se distinto: [[NOME]], CPF [[CPF]], nascimento [[DATA]], relacao [[RELACAO]].")
    add_clause(doc, "1", "Objeto e documentos integrantes", ["A CONTRATADA prestara os servicos odontologicos descritos no Plano de Tratamento e Orcamento aprovados, anexos a este contrato. O plano indicara, por dente/regiao, tecnica, etapas, materiais relevantes, alternativas, profissional responsavel e itens aceitos ou recusados.", "TCLEs, evolucoes, orientacoes, aditivos e documentos de recusa integram o prontuario, mas nao autorizam procedimento nao descrito ou alteracao material sem nova informacao e aceite."])
    add_clause(doc, "2", "Honorarios e pagamento", ["O valor dos itens aprovados e de R$ [[VALOR]], nas condicoes [[CONDICOES]], conforme cronograma anexo. Multa, juros, indice de atualizacao, estorno e politica de ausencia devem ser revisados pelo advogado e informados com destaque.", "Alteracao clinica superveniente, pedido do paciente ou mudanca de escopo exigira novo plano/orcamento. Valores relativos a etapas nao realizadas serao apurados e devolvidos/compensados conforme a lei; etapas efetivamente executadas permanecem devidas."])
    add_clause(doc, "3", "Prazo", ["O prazo estimado e de [[PRAZO]], sujeito a resposta biologica, necessidade de exames, cooperacao, faltas, intercorrencias e disponibilidade de insumos, tudo registrado. Mudanca relevante sera comunicada e documentada; prazo estimado nao e promessa absoluta."])
    add_clause(doc, "4", "Obrigacoes da Contratada", ["A CONTRATADA se obriga a atuar com diligencia, tecnica compativel, seguranca, materiais adequados, informacao clara, respeito a autonomia, confidencialidade, prontuario completo, continuidade responsavel e encaminhamento quando necessario. Nao ha renuncia a garantias legais nem exclusao de responsabilidade por falha profissional ou defeito do servico."])
    add_clause(doc, "5", "Obrigacoes do Contratante/Paciente", ["O paciente ou responsavel se compromete a fornecer informacoes verdadeiras e atualizadas sobre saude, medicamentos e alergias; realizar exames solicitados; comparecer; seguir orientacoes; comunicar sintomas, intercorrencias, mudancas e desistencias; e efetuar pagamentos devidos. Duvidas devem ser apresentadas antes e durante o tratamento."])
    add_clause(doc, "6", "Autonomia, consentimento e segunda opiniao", ["O paciente participara das decisoes, recebera informacao acessivel sobre condicao, alternativas, riscos, beneficios, custos e efeitos de nao tratar, e tera tempo razoavel para decidir, salvo urgencia. Pode buscar segunda opiniao e retirar consentimento antes do procedimento sem represalia, respondendo apenas por custos licitos e comprovados ja incorridos."])
    add_clause(doc, "7", "Natureza do resultado", ["A Odontologia envolve variaveis biologicas e cooperacao. A CONTRATADA nao promete resultado impossivel de garantir; compromete-se com os deveres legais, tecnicos e eticos. Esta clausula nao afasta responsabilidade por erro, defeito, omissao ou informacao insuficiente."])
    add_clause(doc, "8", "Agenda, faltas e abandono", ["Faltas e atrasos serao tratados pela politica [[POLITICA]], previamente informada e proporcional. Ausencia nao caracteriza automaticamente abandono. Antes de encerrar tratamento por falta de continuidade, a CONTRATADA registrara tentativas razoaveis de contato, riscos da interrupcao, prazo para retorno, medidas urgentes e disponibilidade de prontuario/encaminhamento."])
    add_clause(doc, "9", "Rescisao e continuidade segura", ["Qualquer parte pode encerrar o contrato, preservados deveres ja vencidos. A CONTRATADA nao interrompera atendimento de modo a criar risco evitavel: informara situacao clinica, providencias imediatas, documentos, encaminhamento e prazo de transicao compativel. O paciente recebera copia dos documentos solicitados sem exigencia de justificativa."])
    add_clause(doc, "10", "Dados, prontuario e comunicacoes", ["Dados serao tratados conforme o Aviso de Privacidade. O paciente tera acesso e copia gratuita do prontuario nas hipoteses legais. Comunicacoes por WhatsApp/e-mail dependem de contato correto e nao substituem atendimento emergencial. Documentos enviados por canal escolhido devem observar confidencialidade."])
    add_clause(doc, "11", "Forca maior e casos supervenientes", ["Eventos externos inevitaveis podem alterar agenda ou prazo. A parte afetada comunicara e mitigara impactos. Forca maior nao afasta deveres de seguranca, sigilo, informacao, preservacao documental ou devolucao de valores de etapas nao realizadas quando devida."])
    add_clause(doc, "12", "Solucao de conflitos", ["As partes priorizarao esclarecimento pelo canal [[CANAL]], sem impedir acesso ao CRO, autoridades, plataformas de reclamacao ou Judiciario. Aplica-se a lei brasileira e o foro legalmente competente, respeitados direitos do paciente."])
    add_signature_block(doc, ["Paciente/Responsavel Legal: [[NOME/CPF]]", "Cirurgiao-dentista/Clinica: [[NOME/CRO OU CNPJ/EPAO]]"], include_witnesses=True)


def build_budget(doc):
    add_instrument_header(doc, 6, "Plano de Tratamento e Aceite de Orcamento", "Documento de escolha clinica e financeira, vinculado ao contrato")
    add_para(doc, "Clinica: [[NOME/CNPJ/EPAO]]    Profissional: [[NOME/CRO-UF]]")
    add_para(doc, "Paciente: [[NOME/CPF]]    Responsavel: [[NOME/CPF/RELACAO OU NAO SE APLICA]]")
    add_para(doc, "Diagnostico/hipotese e exames de suporte: [[DESCRICAO PROFISSIONAL]]")
    add_heading(doc, "Alternativas discutidas", 2)
    add_table(doc, ["Alternativa", "Beneficios/objetivo", "Limitacoes e riscos", "Escolha"], [
        ["[[ALTERNATIVA 1]]", "[[DESCREVER]]", "[[DESCREVER]]", "[[ACEITA/RECUSA]]"],
        ["[[ALTERNATIVA 2]]", "[[DESCREVER]]", "[[DESCREVER]]", "[[ACEITA/RECUSA]]"],
        ["Nao realizar agora", "Sem intervencao imediata", "[[CONSEQUENCIAS CLINICAS DA RECUSA/ADIAMENTO]]", "[[ACEITA/RECUSA]]"],
    ], [1800, 2400, 2900, 2260])
    add_heading(doc, "Itens aprovados", 2)
    add_table(doc, ["Dente/regiao", "Procedimento/tecnica/material", "Qtd.", "Valor", "Aprovado"], [
        ["[[REGIAO]]", "[[DESCRICAO COMPLETA]]", "[[QTD]]", "R$ [[VALOR]]", "[[SIM/NAO]]"],
        ["[[REGIAO]]", "[[DESCRICAO COMPLETA]]", "[[QTD]]", "R$ [[VALOR]]", "[[SIM/NAO]]"],
    ], [1500, 3900, 800, 1400, 1760])
    add_para(doc, "Total exclusivamente dos itens aprovados: R$ [[TOTAL]]. Validade comercial do orcamento: [[DATA/PRAZO]]. Condicoes: [[ENTRADA/PARCELAS/VENCIMENTOS/MEIO]].")
    add_para(doc, "Itens nao marcados como aprovados permanecem propostos, nao autorizados e nao devidos. Qualquer novo procedimento, mudanca relevante de tecnica/material ou aumento de valor exige informacao e aceite adicional, salvo medida urgente estritamente necessaria para proteger vida/integridade quando o paciente nao puder manifestar vontade, com registro detalhado.")
    add_callout(doc, "DECLARACAO", "Recebi explicacoes acessiveis sobre o plano, alternativas, riscos, beneficios, custos, etapas e consequencias de nao tratar. Pude fazer perguntas e aprovo somente os itens assinalados. Sei que este aceite financeiro nao substitui TCLE especifico quando indicado.", fill=LIGHT_BLUE)
    add_signature_block(doc, ["Paciente/Responsavel: [[NOME/CPF]]", "Cirurgiao-dentista: [[NOME/CRO-UF]]"])


def build_tcle(doc):
    add_instrument_header(doc, 7, "TCLE - Nucleo Geral para Tratamento Odontologico", "Deve ser combinado com o modulo do procedimento e individualizado pelo dentista")
    add_callout(doc, "REGRA DE USO", "Nunca emitir este nucleo sozinho para cirurgia ou procedimento sensivel. O dentista deve selecionar e editar o modulo de risco, preencher campos do caso e assinar junto. IA pode ajudar a organizar o texto, mas nao decide riscos nem substitui a conversa.", fill=LIGHT_GOLD, color=RED)
    add_para(doc, "Clinica: [[NOME/CNPJ/EPAO/ENDERECO/CONTATO]]")
    add_para(doc, "Paciente: [[NOME/CPF/NASCIMENTO]] | Responsavel/representante: [[NOME/CPF/RELACAO/PODERES OU NAO SE APLICA]]")
    add_para(doc, "Cirurgiao-dentista: [[NOME/CRO-UF]] | Auxiliares/outros profissionais relevantes: [[NOMES/FUNCOES]]")
    add_para(doc, "Procedimento, dente/regiao e lateralidade: [[DESCRICAO INEQUIVOCA]]")
    add_clause(doc, "1", "Condicao e indicacao", ["O profissional explicou, em linguagem compreensivel, a condicao observada, os exames relevantes, a indicacao, os objetivos esperados e o prognostico individual: [[DESCRICAO]]."])
    add_clause(doc, "2", "Como sera realizado", ["Tecnica, etapas, anestesia, materiais/dispositivos, duracao estimada e profissionais: [[DESCRICAO]]. Alteracao nao urgente e materialmente diferente dependera de nova conversa e consentimento."])
    add_clause(doc, "3", "Beneficios e expectativas realistas", ["Beneficios esperados e limites: [[DESCRICAO]]. Fatores individuais podem alterar tempo, cicatrizacao e resultado. Nao foi prometido resultado certo; isso nao reduz o dever de diligencia."])
    add_clause(doc, "4", "Alternativas, inclusive nao tratar", ["Foram discutidas as alternativas [[LISTAR]], seus riscos/beneficios/custos e a possibilidade de nao realizar ou adiar. Consequencias de recusa/adiamento: [[DESCRICAO INDIVIDUAL]]."])
    add_clause(doc, "5", "Riscos e complicacoes relevantes", ["Riscos gerais e especificos selecionados no modulo anexo: [[LISTAR E EXPLICAR PROBABILIDADE/GRAVIDADE QUANDO CONHECIDAS]]. Fatores pessoais que aumentam risco: [[TABAGISMO, MEDICAMENTOS, DOENCAS, ANATOMIA, HIGIENE, OUTROS]]. Conduta prevista se ocorrerem: [[DESCRICAO]]."])
    add_clause(doc, "6", "Cuidados e participacao", ["Orientacoes pre e pos-procedimento, medicacao, alimentacao, higiene, restricoes, acompanhante e retorno: [[DESCRICAO]]. O paciente informara mudancas de saude e procurara atendimento diante de [[SINAIS DE ALERTA/CANAL]]."])
    add_clause(doc, "7", "Autonomia, tempo e duvidas", ["Declaro que tive tempo suficiente para decidir, salvo urgencia registrada; pude fazer perguntas; recebi respostas; sei que posso buscar segunda opiniao; e fui informado em formato acessivel. Apoio/interprete utilizado: [[NAO/QUEM/COMO]]."])
    add_clause(doc, "8", "Retirada e limites do consentimento", ["Posso retirar o consentimento antes do inicio, sem represalia. Se o procedimento ja tiver comecado, o profissional explicara riscos de interromper e adotara medidas seguras. A retirada nao apaga registros clinicos nem custos licitos ja incorridos.", "Autorizo apenas medidas adicionais urgentes e estritamente necessarias para proteger minha vida ou integridade caso eu fique incapaz de decidir e nao seja possivel contatar meu representante. Outras ampliacoes exigem novo consentimento."])
    add_clause(doc, "9", "Declaracao final", ["Compreendi a informacao individualizada deste documento e do modulo anexo, recebi oportunidade real de escolha e consinto livremente com o procedimento descrito. Sei que o termo registra uma conversa e nao representa renuncia a direitos ou aceitacao de negligencia."])
    add_signature_block(doc, ["Paciente/Responsavel: [[NOME/CPF/QUALIDADE]]", "Cirurgiao-dentista: [[NOME/CRO-UF]]", "Testemunha da explicacao, se aplicavel: [[NOME/CPF/FUNCAO]]"])


def build_risk_modules(doc):
    add_instrument_header(doc, 8, "Modulos de Risco por Procedimento", "Biblioteca preliminar; selecao e validacao clinica obrigatorias")
    add_callout(doc, "IMPORTANTE", "As listas sao pontos de partida, nao exaustivas nem adequadas a todo caso. Cada especialidade deve validar conteudo, linguagem, frequencia quando disponivel, fatores individuais e sinais de alarme. Itens irrelevantes devem ser removidos; riscos materiais ausentes devem ser acrescentados.", fill=LIGHT_GOLD, color=RED)

    modules = [
        ("A", "Cirurgia oral e extracoes", [
            "Procedimento exato, dente/regiao, motivo, exames e relacao com nervos/seio maxilar.",
            "Dor, inchaco, hematoma, sangramento, limitacao de abertura da boca, infeccao e alveolite.",
            "Lesao de dentes, restauracoes ou estruturas vizinhas; fratura de raiz ou osso; permanencia intencional de fragmento quando remover for mais arriscado.",
            "Comunicacao com seio maxilar, sinusite e necessidade de medidas ou cirurgia adicional quando anatomicamente aplicavel.",
            "Alteracao temporaria ou permanente de sensibilidade, formigamento, dor neuropatica, paladar ou movimento quando houver proximidade nervosa.",
            "Necessidade de ampliar acesso, encaminhar, internar ou realizar procedimento adicional diante de complicacao.",
            "Orientacoes de sangramento, higiene, dieta, medicacao, tabagismo, atividade, retorno e sinais de urgencia.",
        ]),
        ("B", "Implante, enxerto osseo e levantamento de seio", [
            "Local, sistema/material, necessidade de enxerto, exames e planejamento protetico.",
            "Falha de osseointegracao, mobilidade, perda do implante, infeccao, exposicao/reabsorcao do enxerto e necessidade de remocao ou nova cirurgia.",
            "Perda ossea, recessao gengival, alteracao estetica/fonetica, dificuldade de higiene e doenca peri-implantar.",
            "Lesao de dentes, nervos, vasos ou seio maxilar; alteracao sensorial e sinusite quando aplicavel.",
            "Complicacoes proteticas: afrouxamento, fratura, desgaste, desadaptacao e necessidade de manutencao/substituicao.",
            "Impacto de tabagismo, diabetes, higiene, bruxismo, medicamentos e ausencia de manutencao.",
        ]),
        ("C", "Tratamento endodontico e retratamento", [
            "Diagnostico, dente, prognostico, restauracao definitiva necessaria e alternativas (observacao quando cabivel, cirurgia, extracao).",
            "Dor, sensibilidade, inchaco, infeccao, drenagem e necessidade de medicacao ou atendimento adicional.",
            "Anatomia nao localizada, calcificacao, degrau, perfuracao, fratura/separacao de instrumento e extrusao de material.",
            "Persistencia ou retorno da infeccao, escurecimento, fratura do dente/raiz e perda do elemento.",
            "Necessidade de novas sessoes, retratamento, cirurgia apical, encaminhamento ou extracao.",
            "Risco aumentado por trinca, perda estrutural, lesao extensa, reabsorcao, proximidade anatomica ou restauracao inadequada.",
        ]),
        ("D", "Cirurgia e terapia periodontal", [
            "Objetivo, regioes, tecnica, enxerto/material e necessidade de manutencao continuada.",
            "Dor, sangramento, edema, infeccao, sensibilidade e cicatrizacao alterada.",
            "Recessao gengival, exposicao radicular, aumento aparente dos dentes, espacos/triangulos, mudanca estetica e sensibilidade.",
            "Mobilidade, perda de insercao, carie radicular, resultado parcial, recidiva e perda dentaria.",
            "Falha, exposicao ou reabsorcao de enxerto/membrana e necessidade de nova intervencao.",
            "Influencia decisiva de higiene, tabagismo, diabetes, comparecimento e manutencao periodontal.",
        ]),
        ("E", "Ortodontia e alinhadores", [
            "Objetivos, aparelho, duracao estimada, necessidade de exames, desgastes, extracoes, cirurgia ou contencao.",
            "Dor/desconforto, ulceracoes, quebra/descolamento, reacoes a materiais e risco de engolir/aspirar pequenas pecas.",
            "Desmineralizacao, manchas, carie, gengivite, perda periodontal e recessao, especialmente com higiene insuficiente.",
            "Reabsorcao radicular, alteracao pulpar, mobilidade, anquilose ou movimento limitado por fatores biologicos.",
            "Recidiva e necessidade de contencao prolongada; resultado e prazo dependem de cooperacao e resposta individual.",
            "Possivel necessidade de alterar plano, encaminhar ou interromper por condicao periodontal, articular ou falta de cooperacao.",
        ]),
        ("F", "Protese, facetas, coroas e procedimentos esteticos", [
            "Objetivo funcional/estetico, simulacao quando houver, limites de cor, forma, translucidez, simetria e tecidos.",
            "Desgaste irreversivel de estrutura dental; sensibilidade; inflamacão pulpar e eventual necessidade de endodontia.",
            "Fratura, lascamento, descolamento, desgaste, pigmentacao, infiltracao, carie e necessidade de reparo/substituicao.",
            "Alteracoes gengivais/periodontais, dificuldade de higiene, mudancas foneticas e periodo de adaptacao.",
            "Ajustes oclusais, influencia de bruxismo e necessidade de placa/manutencao.",
            "Procedimentos provisórios e etapas laboratoriais podem exigir ajustes de cor/forma antes da instalacao definitiva.",
        ]),
        ("G", "Anestesia local; sedacao exige termo proprio", [
            "Para anestesia local: dor, hematoma, sangramento, infeccao, trismo, desmaio, palpitacao, reacao alergica/toxica, falha anestesica e alteracao sensorial temporaria ou raramente persistente.",
            "Risco de morder/traumatizar labios, lingua ou bochecha enquanto houver dormencia; orientacoes a criancas e responsaveis.",
            "Registrar anestesico, vasoconstritor, dose, lote quando aplicavel, alergias, condicoes e intercorrencias.",
            "Sedacao consciente ou profunda NAO deve usar apenas este modulo. Exige termo especifico para tecnica, profissional habilitado, monitorizacao, jejum, acompanhante, alta, medicacoes e riscos cardiorrespiratorios.",
        ]),
    ]
    for letter, title, items in modules:
        add_heading(doc, f"Modulo {letter} - {title}", 1)
        for item in items:
            add_list_item(doc, item)
        add_para(doc, "Fatores individuais selecionados: [[PREENCHER]]. Riscos explicados ao paciente: [[PREENCHER]]. Conduta/sinais de alarme: [[PREENCHER]].", italic=True, color=DARK_BLUE)


def build_refusal(doc):
    add_instrument_header(doc, 9, "Termo de Recusa Informada, Interrupcao ou Descontinuidade", "Para registrar decisao do paciente sem transformar recusa em abandono automatico")
    add_para(doc, "Paciente/Responsavel: [[NOME/CPF/RELACAO]] | Profissional: [[NOME/CRO-UF]] | Data/hora: [[DATA/HORA]]")
    add_clause(doc, "1", "Recomendacao profissional", ["Condicao e achados: [[DESCRICAO]]. Tratamento/exame/retorno recomendado: [[DESCRICAO]]. Urgencia e prazo clinico: [[DESCRICAO]]."])
    add_clause(doc, "2", "Informacao prestada", ["Foram explicados objetivo, beneficios, riscos, alternativas, custos e efeitos provaveis de adiar ou recusar: [[DESCRICAO INDIVIDUAL]]. O paciente recebeu oportunidade para perguntas e segunda opiniao."])
    add_clause(doc, "3", "Decisao do paciente", ["O paciente/responsavel escolhe: [[RECUSAR / ADIAR ATE DATA / INTERROMPER / TRANSFERIR]]. Motivo declarado, se desejar registrar: [[MOTIVO OU 'PREFERE NAO INFORMAR']]. A recusa e livre e nao gera represalia."])
    add_clause(doc, "4", "Plano de seguranca e continuidade", ["Sinais de alarme e onde buscar atendimento: [[DESCRICAO]]. Medidas temporarias: [[DESCRICAO]]. Retorno oferecido: [[DATA/CANAL]]. Encaminhamento/copia de documentos: [[DESCRICAO]]."])
    add_clause(doc, "5", "Registro de contato", ["Quando o paciente deixa de comparecer, registrar tentativas razoaveis e proporcionais de contato, sem expor dados de saude a terceiros: [[DATAS/CANAIS/RESULTADOS]]. Nao usar este termo para fabricar assinatura ou presumir consentimento."])
    add_callout(doc, "DECLARACAO", "Compreendi a recomendacao e os riscos explicados e, neste momento, mantenho a decisao acima. Sei que posso procurar atendimento ou retomar contato. Este registro nao afasta o dever do profissional de agir com seguranca nem meus direitos legais.", fill=LIGHT_BLUE)
    add_signature_block(doc, ["Paciente/Responsavel: [[NOME/CPF]]", "Cirurgiao-dentista: [[NOME/CRO-UF]]", "Testemunha da recusa verbal, se o paciente nao assinar: [[NOME/CPF/FUNCAO]]"])


def build_completion(doc):
    add_instrument_header(doc, 10, "Registro de Conclusao, Intercorrencias e Orientacoes", "Documento pos-procedimento vinculado a evolucao clinica; nao e quitacao geral")
    add_para(doc, "Paciente: [[NOME/CPF]] | Profissional executor: [[NOME/CRO-UF]] | Sessao: [[DATA/HORA]]")
    add_heading(doc, "Procedimentos efetivamente realizados", 2)
    add_table(doc, ["Dente/regiao", "Procedimento/tecnica", "Material/lote quando relevante", "Executor"], [
        ["[[REGIAO]]", "[[DESCRICAO]]", "[[MATERIAL/LOTE/NA]]", "[[NOME/CRO]]"],
    ], [1500, 3300, 2700, 1860])
    add_heading(doc, "Intercorrencias e resposta", 2)
    add_para(doc, "[[DESCREVER FATOS, CONDUTAS E RESULTADOS, OU 'NAO HOUVE INTERCORRENCIA REGISTRADA']]")
    add_heading(doc, "Orientacoes entregues", 2)
    for item in [
        "Medicacao e uso: [[DESCREVER OU NAO SE APLICA]].",
        "Higiene, dieta, atividade, tabagismo/alcool e restricoes: [[DESCREVER]].",
        "Sinais esperados e sinais de alarme: [[DESCREVER]].",
        "Canal e local para urgencia: [[CONTATO/ENDERECO]].",
        "Retorno: [[DATA/PRAZO/OBJETIVO]].",
        "Documento/copia entregue por: [[IMPRESSO/E-MAIL/PORTAL/WHATSAPP AUTORIZADO]].",
    ]:
        add_list_item(doc, item)
    add_callout(doc, "CIENCIA", "Confirmo que recebi as orientacoes acima em linguagem compreensivel, pude perguntar e sei como buscar ajuda. Minha assinatura registra recebimento e fatos desta sessao; nao significa renuncia a direitos, garantia de resultado ou quitacao ampla.", fill=LIGHT_BLUE)
    add_signature_block(doc, ["Paciente/Responsavel: [[NOME/CPF]]", "Cirurgiao-dentista executor: [[NOME/CRO-UF]]"])


def build_recording(doc):
    add_instrument_header(doc, 11, "Autorizacao Especifica para Audio, Voz, Imagem e Transcricao", "Separar registro clinico, ensino/pesquisa e marketing; uma finalidade por opcao")
    add_callout(doc, "DESENHO DE PRODUTO", "Preferencia de menor risco: modo ditado em que apenas o dentista fala, fora da conversa do paciente. Se a voz do paciente puder ser capturada no ambiente, mostrar aviso e coletar escolha especifica antes de gravar, oferecendo fluxo sem gravacao.", fill=LIGHT_BLUE)
    add_para(doc, "Paciente/Responsavel: [[NOME/CPF/RELACAO]] | Clinica: [[NOME/CNPJ]] | Data: [[DATA]]")
    add_para(doc, "Marque separadamente:")
    add_list_item(doc, "[[  ]] Autorizo gravacao de audio da consulta para documentacao clinica e transcricao, com revisao do cirurgiao-dentista e guarda no prontuario por [[PRAZO/FORMA]].")
    add_list_item(doc, "[[  ]] Autorizo captura de fotografias/imagens clinicas para diagnostico, planejamento, acompanhamento e prontuario, sem publicacao.")
    add_list_item(doc, "[[  ]] Autorizo uso anonimizado para ensino/cientifico, desde que nao seja razoavelmente possivel me identificar. Condicoes: [[DETALHAR]].")
    add_list_item(doc, "[[  ]] Autorizo uso identificavel em divulgacao/marketing nos canais [[LISTAR]], pelo prazo [[PRAZO]] e territorio [[TERRITORIO]]. Esta opcao e facultativa e deve permanecer desmarcada por padrao.")
    add_para(doc, "Finalidade, destinatarios, fornecedores, paises e periodo de retencao: [[DETALHAR DE FORMA ESPECIFICA]].")
    add_para(doc, "Posso recusar opcoes facultativas sem prejuizo do atendimento e revogar autorizacao para usos futuros pelo canal [[CANAL]]. A revogacao nao torna ilicito o uso anterior nem exige apagar registro clinico cuja guarda seja necessaria, mas interrompera novas publicacoes/usos quando juridica e tecnicamente possivel.")
    add_para(doc, "A Clinica nao vendera a gravacao nem a utilizara para treinar modelos de proposito geral sem novo instrumento especifico. O cirurgiao-dentista revisara a transcricao antes de integra-la ao prontuario.")
    add_signature_block(doc, ["Paciente/Responsavel: [[NOME/CPF]]", "Responsavel pela coleta: [[NOME/CRO OU FUNCAO]]"])


def build_operational_rules(doc):
    add_instrument_header(doc, 12, "Politica Interna de Geracao e Assinatura", "Checklist de produto e operacao; nao exibido integralmente ao paciente")
    add_heading(doc, "1. Regras de geracao", 1)
    for item in [
        "Documento nasce de dados remontados no servidor e da clinica ativa; navegador nao e fonte de valores, CRO, paciente, procedimento ou autoria.",
        "Template possui versao semantica, data de vigencia, aprovador juridico, aprovador clinico, especialidade, hash e historico.",
        "Campos clinicos obrigatorios nao aceitam texto generico como 'riscos habituais' ou 'conforme explicado'.",
        "IA nunca preenche automaticamente diagnostico, risco aplicavel ou declaracao de que houve explicacao. Ela pode sugerir texto marcado como rascunho.",
        "Documento final mostra ao profissional e ao paciente uma previa completa antes da assinatura.",
        "Alteracao posterior gera adendo; o original e o log permanecem imutaveis.",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "2. Escada de identidade e assinatura", 1)
    add_table(doc, ["Nivel", "Uso sugerido", "Evidencia minima"], [
        ["Basico", "Termos de usuario e documentos de baixo risco", "Sessao autenticada, checkbox destacado, versao/hash, data/hora, IP, copia"],
        ["Reforcado", "Orcamento e TCLE comum", "Identificacao do paciente, assinatura desenhada ou OTP, assinatura do dentista, trilha e PDF imutavel"],
        ["Alto risco", "Cirurgia complexa, sequela possivel, alto custo", "Assinatura avancada/qualificada conforme parecer; testemunha ou registro adicional; dialogo documentado"],
    ], [1500, 3100, 4760])
    add_heading(doc, "3. Gates antes de publicar", 1)
    for item in [
        "Razao social, CNPJ, endereco, suporte, privacidade e foro preenchidos.",
        "Advogado aprovou cada minuta e a relacao entre Terms, DPA, Avisos e contratos clinicos.",
        "Diretor clinico e especialistas validaram modulos de risco.",
        "Suboperadores, paises, retencao, treinamento de IA e transferencias internacionais confirmados por contrato.",
        "Politica real de backup, exportacao, exclusao, RTO/RPO e incidentes confere com o texto.",
        "Fluxo diferencia titular/admin de membro convidado e registra poderes de representacao.",
        "Teste com duas contas e duas clinicas confirma isolamento; PDFs e snapshots permanecem inteiros apos mudanca de dados.",
        "Usuario recebe copia; nova versao material pede novo aceite; recusas e revogacoes ficam registradas.",
        "Auditoria de acessibilidade e linguagem simples em celular, incluindo leitor de tela e paciente com apoio/interprete.",
        "Feature flag permanece desligada ate aprovacao juridica e operacional.",
    ]:
        add_list_item(doc, item)


def build_sources(doc):
    heading = add_heading(doc, "Referencias oficiais para revisao juridica", 1)
    heading.paragraph_format.page_break_before = True
    add_para(doc, "As minutas foram estruturadas a partir destas fontes. A citacao nao substitui interpretacao profissional nem confirma adequacao a um caso concreto.")
    sources = [
        ("Estatuto dos Direitos do Paciente - Lei 15.378/2026", "https://planalto.gov.br/ccivil_03/_ato2023-2026/2026/lei/l15378.htm"),
        ("Manual do Prontuario do Paciente em Odontologia - CFO, 2026", "https://website.cfo.org.br/wp-content/uploads/2026/03/CFO_Manual_do_Prontuario_Ebook_v2.pdf"),
        ("Codigo de Etica Odontologica - Resolucao CFO 118/2012", "https://website.cfo.org.br/wp-content/uploads/2018/03/codigo_etica.pdf"),
        ("Lei Geral de Protecao de Dados - Lei 13.709/2018", "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709compilado.htm"),
        ("Guarda e digitalizacao de prontuario - Lei 13.787/2018", "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13787.htm"),
        ("Documentos e assinaturas eletronicas - MP 2.200-2/2001", "https://www.planalto.gov.br/ccivil_03/mpv/antigas_2001/2200-2.htm"),
        ("Assinaturas eletronicas - Lei 14.063/2020", "https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2020/lei/l14063.htm"),
        ("Marco Civil da Internet - Lei 12.965/2014", "https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2014/lei/l12965.htm"),
        ("Codigo de Defesa do Consumidor - Lei 8.078/1990", "https://www.planalto.gov.br/ccivil_03/leis/l8078compilado.htm"),
        ("ANPD - Comunicacao de incidente, Resolucao 15/2024", "https://www.gov.br/anpd/pt-br/canais_atendimento/agente-de-tratamento/comunicado-de-incidente-de-seguranca-cis"),
        ("ANPD - Transferencia internacional, Resolucao 19/2024", "https://www.gov.br/anpd/pt-br/acesso-a-informacao/institucional/atos-normativos/regulamentacoes_anpd/resolucao-cd-anpd-no-19-de-23-de-agosto-de-2024"),
        ("ANPD - Guia de agentes de tratamento e encarregado", "https://www.gov.br/anpd/pt-br/centrais-de-conteudo/materiais-educativos-e-publicacoes/guia-orientativo-para-definicoes-dos-agentes-de-tratamento-de-dados-pessoais-e-do-encarregado"),
        ("AAE - Informed Consent (apoio clinico, adaptar ao Brasil)", "https://www.aae.org/specialty/clinical-resources/treatment-planning/informed-consent/"),
    ]
    for label, url in sources:
        p = add_list_item(doc, "")
        add_runs(p, f"{label}: ", bold=True, color=INK)
        add_runs(p, url, color=BLUE, size=9.5)
    add_heading(doc, "Nota final ao advogado", 1)
    add_para(doc, "Revisar especialmente: qualificacao das partes; incidencia do CDC; limite de responsabilidade; foro; cobranca/renovacao; assinatura e executividade; bases legais por finalidade; papeis controlador-operador; transferencias; subprocessadores; retencao; prontuario de menor/incapaz; audio ambiente; e compatibilidade dos modulos clinicos com cada especialidade e caso.")


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    build_front_matter(doc)
    build_terms(doc)
    build_user_privacy(doc)
    build_dpa(doc)
    build_patient_privacy(doc)
    build_service_contract(doc)
    build_budget(doc)
    build_tcle(doc)
    build_risk_modules(doc)
    build_refusal(doc)
    build_completion(doc)
    build_recording(doc)
    build_operational_rules(doc)
    build_sources(doc)
    setup_header_footer(doc)
    doc.core_properties.title = "Dossie juridico preliminar - Odonto.IA"
    doc.core_properties.subject = "Minutas para revisao juridica e validacao clinica"
    doc.core_properties.author = "Odonto.IA"
    doc.core_properties.keywords = "odontologia, contratos, TCLE, LGPD, SaaS, prontuario"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
