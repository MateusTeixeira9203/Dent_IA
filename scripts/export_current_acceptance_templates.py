from __future__ import annotations

from pathlib import Path

from docx import Document

import build_legal_dossier as base
from build_separate_legal_documents import finish_document


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "juridico" / "documentos-separados"
TEMPLATE_VERSION = "1.0-draft"


def add_current_header(doc: Document, title: str, source_function: str) -> None:
    base.add_para(doc, "MODELO ATUAL DO SISTEMA", bold=True, color=base.GOLD, size=10, after=4, keep=True)
    base.add_para(doc, title.upper(), bold=True, color=base.INK, size=20, after=4, keep=True)
    base.add_para(
        doc,
        f"Transcricao para revisao juridica | Template {TEMPLATE_VERSION} | Fonte: src/lib/legal/templates.ts ({source_function})",
        italic=True,
        color=base.MUTED,
        size=10.5,
        after=12,
    )
    base.add_callout(
        doc,
        "STATUS",
        "Este arquivo reproduz o modelo que o sistema gera hoje. O proprio codigo o classifica como draft. Nao confundir com as minutas juridicas ampliadas do pacote e nao ativar em producao sem revisao do advogado.",
        fill=base.LIGHT_GOLD,
        color=base.RED,
    )


def add_pdf_layer_note(doc: Document) -> None:
    base.add_heading(doc, "Camada acrescentada pelo PDF atual", 2)
    base.add_para(
        doc,
        "Fora do texto do template, o gerador inclui novamente paciente/CPF no cabecalho, data, imagem da assinatura do paciente ou responsavel, nome de quem assinou e uma linha com nome e CRO do dentista. No fluxo atual de aceite, nao e enviada imagem de assinatura do dentista ao PDF.",
        color=base.MUTED,
        size=10,
    )


def add_identity(doc: Document, *, executor: bool = False) -> None:
    base.add_para(doc, "Clinica: [[NOME DA CLINICA]]")
    base.add_para(doc, "Paciente: [[NOME DO PACIENTE]] - CPF [[CPF OU AUSENTE]]")
    base.add_para(doc, "Representante legal: [[NOME/CPF OU LINHA AUSENTE]]")
    role = "Cirurgiao-dentista executor" if executor else "Cirurgiao-dentista"
    base.add_para(doc, f"{role}: [[NOME DO DENTISTA]] - CRO [[CRO OU AUSENTE]]", after=10)


def build_tcle() -> Path:
    doc = Document()
    base.setup_styles(doc)
    add_current_header(doc, "TCLE - aceite pre-procedimento", "renderTCLE")
    base.add_heading(doc, "Texto atual do template", 1)
    base.add_para(doc, "TERMO DE CONSENTIMENTO LIVRE E ESCLARECIDO - PROCEDIMENTO ODONTOLOGICO", bold=True, color=base.INK, size=12)
    add_identity(doc)
    sections = [
        ("PROCEDIMENTO", "[[PROCEDIMENTO SELECIONADO NO PRONTUARIO]]"),
        ("REGIAO", "[[DENTE OU REGIAO CLINICA REGISTRADA]]"),
        ("POR QUE FOI INDICADO", "[[JUSTIFICATIVA PREENCHIDA PELO DENTISTA]]"),
        ("COMO SERA REALIZADO", "[[EXPLICACAO PREENCHIDA PELO DENTISTA]]"),
        ("ALTERNATIVAS APRESENTADAS", "[[ALTERNATIVAS PREENCHIDAS PELO DENTISTA]]"),
        ("CONSEQUENCIAS DE NAO REALIZAR", "[[CONSEQUENCIAS PREENCHIDAS PELO DENTISTA]]"),
        ("RISCOS E COMPLICACOES POSSIVEIS", "[[RISCOS PREENCHIDOS PELO DENTISTA]]"),
        ("ORIENTACOES E CUIDADOS", "[[ORIENTACOES PREENCHIDAS PELO DENTISTA]]"),
    ]
    for label, value in sections:
        base.add_para(doc, label, bold=True, color=base.DARK_BLUE, size=10.5, after=2, keep=True)
        base.add_para(doc, value, after=6)
    base.add_callout(
        doc,
        "DECLARACAO ATUAL",
        "Declaro que recebi explicacoes compreensiveis, pude tirar duvidas e consinto livremente com a realizacao do procedimento descrito. Estou ciente de que posso revogar este consentimento antes do inicio do procedimento, observado o registro clinico e eventuais custos comprovadamente ja incorridos.",
        fill=base.LIGHT_BLUE,
    )
    add_pdf_layer_note(doc)
    path = OUTPUT_DIR / "13-modelo-atual-tcle-aceite-pre-procedimento.docx"
    finish_document(doc, "Modelo atual - TCLE e aceite pre-procedimento", path)
    return path


def build_completion() -> Path:
    doc = Document()
    base.setup_styles(doc)
    add_current_header(doc, "Aceite de conclusao do procedimento realizado", "renderConclusao")
    base.add_heading(doc, "Texto atual do template", 1)
    base.add_para(doc, "TERMO DE CONCLUSAO E ACEITE DE PROCEDIMENTO REALIZADO", bold=True, color=base.INK, size=12)
    add_identity(doc, executor=True)
    base.add_para(doc, "PROCEDIMENTOS REALIZADOS NESTA SESSAO", bold=True, color=base.DARK_BLUE, size=10.5, after=2, keep=True)
    base.add_para(doc, "1. [[PROCEDIMENTO, DENTE E OBSERVACAO REGISTRADOS]]", after=6)
    base.add_para(doc, "ORIENTACOES ENTREGUES", bold=True, color=base.DARK_BLUE, size=10.5, after=2, keep=True)
    base.add_para(doc, "[[ORIENTACOES PREENCHIDAS NO FLUXO DE ASSINATURA]]", after=6)
    base.add_para(doc, "INTERCORRENCIAS", bold=True, color=base.DARK_BLUE, size=10.5, after=2, keep=True)
    base.add_para(doc, "[[INTERCORRENCIA INFORMADA OU 'Nao houve intercorrencia registrada.']]", after=6)
    base.add_para(doc, "RETORNO", bold=True, color=base.DARK_BLUE, size=10.5, after=2, keep=True)
    base.add_para(doc, "[[RETORNO INFORMADO OU 'Conforme orientacao clinica.']]", after=6)
    base.add_callout(
        doc,
        "DECLARACAO ATUAL",
        "Declaro que os procedimentos acima foram realizados nesta sessao, que recebi as orientacoes descritas e que pude esclarecer duvidas. Este aceite comprova a execucao e a entrega das orientacoes, sem limitar garantias ou direitos legais aplicaveis.",
        fill=base.LIGHT_BLUE,
    )
    add_pdf_layer_note(doc)
    path = OUTPUT_DIR / "14-modelo-atual-aceite-conclusao-procedimento.docx"
    finish_document(doc, "Modelo atual - aceite de conclusao de procedimento", path)
    return path


def build_budget_acceptance() -> Path:
    doc = Document()
    base.setup_styles(doc)
    add_current_header(doc, "Aceite de orcamento e plano de tratamento", "renderAceiteOrcamento")
    base.add_heading(doc, "Texto atual do template", 1)
    base.add_para(doc, "TERMO DE ACEITE DE ORCAMENTO E PLANO DE TRATAMENTO ODONTOLOGICO", bold=True, color=base.INK, size=12)
    base.add_para(doc, "Clinica: [[NOME DA CLINICA]]")
    base.add_para(doc, "Paciente: [[NOME DO PACIENTE]] - CPF [[CPF OU AUSENTE]]")
    base.add_para(doc, "Profissional responsavel: [[NOME DO DENTISTA]] - CRO [[CRO OU AUSENTE]]", after=10)
    base.add_para(doc, "PROCEDIMENTOS APROVADOS", bold=True, color=base.DARK_BLUE, size=10.5, after=2, keep=True)
    base.add_para(doc, "1. [[QUANTIDADE]] x [[DESCRICAO DO ITEM APROVADO]] - R$ [[PRECO TOTAL DO ITEM]]", after=3)
    base.add_para(doc, "2. [[DEMAIS ITENS APROVADOS; ITENS NAO APROVADOS NAO ENTRAM]]", after=8)
    base.add_para(doc, "VALOR TOTAL APROVADO: R$ [[TOTAL DO SNAPSHOT]]", bold=True, color=base.DARK_BLUE, size=11)
    base.add_para(doc, "Condicoes de pagamento: [[CONDICOES REGISTRADAS OU 'A definir com a clinica.']]", after=10)
    base.add_callout(
        doc,
        "DECLARACAO ATUAL",
        "Declaro que recebi explicacao em linguagem acessivel sobre o plano aprovado, suas alternativas, limites e condicoes financeiras. Aceito exclusivamente os procedimentos acima, ciente de que qualquer alteracao relevante ou novo valor depende de novo orcamento e novo aceite.\n\nO presente documento registra manifestacao de vontade eletronica do paciente ou responsavel no ato da assinatura.",
        fill=base.LIGHT_BLUE,
    )
    add_pdf_layer_note(doc)
    path = OUTPUT_DIR / "15-modelo-atual-aceite-orcamento.docx"
    finish_document(doc, "Modelo atual - aceite de orcamento", path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for output in (build_tcle(), build_completion(), build_budget_acceptance()):
        print(output)


if __name__ == "__main__":
    main()
